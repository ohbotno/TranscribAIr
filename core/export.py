"""
Export functionality for organized feedback.
Supports clipboard, text, PDF, and Word document formats.
"""
from pathlib import Path
from typing import Optional
from datetime import datetime

from core.feedback import OrganizedFeedback


class FeedbackExporter:
    """Handles exporting organized feedback to various formats."""

    @staticmethod
    def to_clipboard(feedback: OrganizedFeedback, format: str = "plain") -> bool:
        """
        Copy feedback to clipboard.

        Args:
            feedback: OrganizedFeedback object
            format: "plain" or "markdown"

        Returns:
            True if successful
        """
        try:
            import tkinter as tk
            root = tk.Tk()
            root.withdraw()

            text = (
                feedback.to_plain_text()
                if format == "plain"
                else feedback.to_markdown()
            )

            root.clipboard_clear()
            root.clipboard_append(text)
            root.update()
            root.destroy()

            return True
        except Exception as e:
            print(f"Error copying to clipboard: {e}")
            return False

    @staticmethod
    def to_text_file(feedback: OrganizedFeedback, file_path: Path) -> bool:
        """
        Export feedback to plain text file.

        Args:
            feedback: OrganizedFeedback object
            file_path: Path to save the file

        Returns:
            True if successful
        """
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(feedback.to_plain_text())
            return True
        except Exception as e:
            print(f"Error exporting to text file: {e}")
            return False

    @staticmethod
    def to_markdown_file(feedback: OrganizedFeedback, file_path: Path) -> bool:
        """
        Export feedback to markdown file.

        Args:
            feedback: OrganizedFeedback object
            file_path: Path to save the file

        Returns:
            True if successful
        """
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(feedback.to_markdown())
            return True
        except Exception as e:
            print(f"Error exporting to markdown file: {e}")
            return False

    @staticmethod
    def to_pdf(feedback: OrganizedFeedback, file_path: Path) -> bool:
        """
        Export feedback to PDF file.

        Args:
            feedback: OrganizedFeedback object
            file_path: Path to save the file

        Returns:
            True if successful
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY

            doc = SimpleDocTemplate(
                str(file_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18,
            )

            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='Justify',
                parent=styles['Normal'],
                alignment=TA_JUSTIFY,
                spaceAfter=12
            ))

            story = []

            # Title
            title = Paragraph(
                f"<b>Feedback: {feedback.rubric_name}</b>",
                styles['Title']
            )
            story.append(title)
            story.append(Spacer(1, 0.2 * inch))

            # Date
            date_text = Paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                styles['Normal']
            )
            story.append(date_text)
            story.append(Spacer(1, 0.3 * inch))

            # Summary
            if feedback.summary:
                summary_heading = Paragraph("<b>Summary</b>", styles['Heading2'])
                story.append(summary_heading)
                summary_text = Paragraph(feedback.summary, styles['Justify'])
                story.append(summary_text)
                story.append(Spacer(1, 0.2 * inch))

            # Detailed Feedback
            feedback_heading = Paragraph("<b>Detailed Feedback</b>", styles['Heading2'])
            story.append(feedback_heading)
            story.append(Spacer(1, 0.1 * inch))

            for criterion, text in feedback.criterion_feedback.items():
                criterion_heading = Paragraph(f"<b>{criterion}</b>", styles['Heading3'])
                story.append(criterion_heading)

                criterion_text = Paragraph(text, styles['Justify'])
                story.append(criterion_text)
                story.append(Spacer(1, 0.15 * inch))

            # Raw Transcript (if included)
            if feedback.raw_transcript:
                story.append(PageBreak())
                transcript_heading = Paragraph("<b>Raw Transcript</b>", styles['Heading2'])
                story.append(transcript_heading)
                story.append(Spacer(1, 0.1 * inch))

                transcript_text = Paragraph(feedback.raw_transcript, styles['Normal'])
                story.append(transcript_text)

            doc.build(story)
            return True

        except ImportError:
            print("reportlab not installed. Install with: pip install reportlab")
            return False
        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            return False

    @staticmethod
    def to_word(feedback: OrganizedFeedback, file_path: Path) -> bool:
        """
        Export feedback to Word document.

        Args:
            feedback: OrganizedFeedback object
            file_path: Path to save the file

        Returns:
            True if successful
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading(f"Feedback: {feedback.rubric_name}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Date
            date_para = doc.add_paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            date_para.runs[0].font.size = Pt(10)

            doc.add_paragraph()  # Spacing

            # Summary
            if feedback.summary:
                doc.add_heading("Summary", level=2)
                summary_para = doc.add_paragraph(feedback.summary)
                summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph()  # Spacing

            # Detailed Feedback
            doc.add_heading("Detailed Feedback", level=2)

            for criterion, text in feedback.criterion_feedback.items():
                doc.add_heading(criterion, level=3)
                feedback_para = doc.add_paragraph(text)
                feedback_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph()  # Spacing

            # Raw Transcript (if included)
            if feedback.raw_transcript:
                doc.add_page_break()
                doc.add_heading("Raw Transcript", level=2)
                transcript_para = doc.add_paragraph(feedback.raw_transcript)

            doc.save(str(file_path))
            return True

        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return False
        except Exception as e:
            print(f"Error exporting to Word: {e}")
            return False

    @staticmethod
    def export(
        feedback: OrganizedFeedback,
        file_path: Optional[Path] = None,
        format: str = "clipboard"
    ) -> bool:
        """
        Export feedback to the specified format.

        Args:
            feedback: OrganizedFeedback object
            file_path: Path to save the file (required for file formats)
            format: "clipboard", "text", "markdown", "pdf", or "word"

        Returns:
            True if successful
        """
        if format == "clipboard":
            return FeedbackExporter.to_clipboard(feedback, "plain")
        elif format == "clipboard_markdown":
            return FeedbackExporter.to_clipboard(feedback, "markdown")
        elif format == "text":
            if not file_path:
                raise ValueError("file_path required for text format")
            return FeedbackExporter.to_text_file(feedback, file_path)
        elif format == "markdown":
            if not file_path:
                raise ValueError("file_path required for markdown format")
            return FeedbackExporter.to_markdown_file(feedback, file_path)
        elif format == "pdf":
            if not file_path:
                raise ValueError("file_path required for PDF format")
            return FeedbackExporter.to_pdf(feedback, file_path)
        elif format == "word":
            if not file_path:
                raise ValueError("file_path required for Word format")
            return FeedbackExporter.to_word(feedback, file_path)
        else:
            raise ValueError(f"Unknown format: {format}")
